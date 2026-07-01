import streamlit as st
import re
import io

# ─── CONFIGURAÇÃO DA PÁGINA ────────────────────────────────────────
st.set_page_config(
    page_title="Gerador Batch — PRODAM",
    page_icon="📊",
    layout="wide",
)

# ─── CSS CUSTOMIZADO ───────────────────────────────────────────────
st.markdown("""
<style>
    .main .block-container { max-width: 1200px; padding-top: 1rem; }
    .header-bar {
        background: #1F4E79; color: #A8C8E8; padding: 12px 20px;
        border-radius: 6px; margin-bottom: 16px;
        display: flex; align-items: center; justify-content: space-between;
    }
    .header-bar h2 { color: #fff; margin: 0; font-size: 18px; }
    .header-bar .badge { background: #2E75B6; color: #A8C8E8; font-size: 11px; padding: 3px 10px; border-radius: 12px; }
    .success-box { background: #E8F5E9; border-left: 3px solid #1B5E20; padding: 10px 14px; border-radius: 4px; margin: 8px 0; }
    .info-box { background: #E3F0FA; border-left: 3px solid #1565A0; padding: 10px 14px; border-radius: 4px; margin: 8px 0; }
    .warn-box { background: #FFF3E0; border-left: 3px solid #E65100; padding: 10px 14px; border-radius: 4px; margin: 8px 0; }
</style>
""", unsafe_allow_html=True)

# ─── CABEÇALHO ─────────────────────────────────────────────────────
st.markdown("""
<div class="header-bar">
    <div>
        <h2>■ PRODAM — Gerador de Documentação Batch</h2>
        <div style="font-size:12px;color:#7AABCC;margin-top:2px">Upload do(s) script(s) SQL → geração automática da Característica (.docx) e do Fluxograma (SVG)</div>
    </div>
    <span class="badge">SIGA SAÚDE · V3.0</span>
</div>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════
# PARSER DE SQL
# ═══════════════════════════════════════════════════════════════════

def extrair_path_do_texto(texto: str) -> str:
    """Acha a primeira linha que parece um caminho de rede (começa com '/') num texto livre."""
    for linha in str(texto or '').split('\n'):
        linha = linha.strip()
        if linha.startswith('/'):
            return linha
    return ''


def extrair_dados_sql(texto: str, nome_arquivo: str = '') -> dict:
    """Extrai informações do script SQL."""
    dados = {
        'tabelas': [],
        'arquivo_saida': '',
        'nome_script': nome_arquivo,
        'tem_select': False,
        'tem_insert': False,
        'tem_update': False,
        'tem_delete': False,
    }

    # Tabelas (FROM e JOIN) — exclui DUAL
    tabelas = set()
    for m in re.finditer(r'(?:FROM|JOIN)\s+(TB_\w+|VW_\w+)', texto, re.IGNORECASE):
        tabelas.add(m.group(1).upper())
    dados['tabelas'] = sorted(tabelas)

    # Arquivo de saída (SPOOL) — ignora SPOOL ON/OFF/OUT
    for m in re.finditer(r'SPOOL\s+(\S+)', texto, re.IGNORECASE):
        val = m.group(1).strip()
        if val.upper() not in ('ON', 'OFF', 'OUT'):
            dados['arquivo_saida'] = val
            break

    # Tipo de operação
    dados['tem_select'] = bool(re.search(r'\bSELECT\b', texto, re.IGNORECASE))
    dados['tem_insert'] = bool(re.search(r'\bINSERT\b', texto, re.IGNORECASE))
    dados['tem_update'] = bool(re.search(r'\bUPDATE\b', texto, re.IGNORECASE))
    dados['tem_delete'] = bool(re.search(r'\bDELETE\b', texto, re.IGNORECASE))

    return dados


def descrever_operacao_sql(dados_sql: dict) -> str:
    """Gera descrição textual do que o script faz."""
    ops = []
    if dados_sql['tem_select']:
        ops.append('consulta (SELECT)')
    if dados_sql['tem_insert']:
        ops.append('inserção (INSERT)')
    if dados_sql['tem_update']:
        ops.append('atualização (UPDATE)')
    if dados_sql['tem_delete']:
        ops.append('exclusão (DELETE)')
    if ops:
        return 'Script faz ' + ', '.join(ops) + ' no banco.'
    return 'Script SQL.'


# ═══════════════════════════════════════════════════════════════════
# GERADOR DE CARACTERÍSTICA (.docx)
# ═══════════════════════════════════════════════════════════════════

def gerar_caracteristica_docx(dados: dict) -> bytes:
    """Gera o documento Característica no formato oficial PRODAM usando python-docx."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial Narrow'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(3.6)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.footer_distance = Cm(0.3)

    PRETO = '000000'

    cod = dados.get('cod', '')
    sis = dados.get('sis', 'SIGA SAUDE')
    sub = dados.get('sub', '')
    den = dados.get('den', '')
    amb = dados.get('amb', 'UNIX')
    dat = dados.get('dat', '')
    hora = dados.get('hora', '00h00')
    freq = dados.get('freq', 'diariamente')
    etapas = dados.get('etapas', [])

    # ─── HELPERS ────────────────────────────────────────────────
    def set_table_borders(tbl, color=PRETO, sz=4):
        """Bordas no nível da TABELA (tblBorders) — em vez de célula a célula.
        Isso evita o bug de células mescladas verticalmente (vMerge) ficarem
        sem borda na 'continuação' da mesclagem, o que fazia a tabela parecer
        deslocada/cortada."""
        tblPr = tbl._tbl.tblPr
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        ))

    def cell_text(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0, underline=False):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.underline = underline
        run.font.name = 'Arial Narrow'
        run.font.color.rgb = RGBColor.from_string(PRETO)
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return p

    def cell_label_value(cell, label, value, size_label=6, size_value=9,
                          align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(str(label))
        r1.font.size = Pt(size_label)
        r1.font.name = 'Arial Narrow'
        r1.font.color.rgb = RGBColor.from_string(PRETO)

        p2 = cell.add_paragraph()
        p2.alignment = align
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(str(value))
        r2.font.size = Pt(size_value)
        r2.font.bold = True
        r2.font.name = 'Arial Narrow'
        r2.font.color.rgb = RGBColor.from_string(PRETO)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def set_col_widths(tbl, widths_cm):
        tbl.autofit = False
        for i, w in enumerate(widths_cm):
            tbl.columns[i].width = Cm(w)
        for row in tbl.rows:
            for cell, w in zip(row.cells, widths_cm):
                cell.width = Cm(w)

    def add_texto(text, size=11, bold=False, underline=False, space_after=9, space_before=0):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.underline = underline
        run.font.name = 'Arial Narrow'
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        return p

    AZUL = '0563C1'

    def is_path(text):
        t = str(text or '').strip()
        return t.startswith('/') or t.startswith('\\\\') or t.startswith('http')

    def add_link_texto(text, size=11, space_after=9, space_before=0):
        """Parágrafo em azul sublinhado — usado para caminhos/pastas de rede."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.underline = True
        run.font.name = 'Arial Narrow'
        run.font.color.rgb = RGBColor.from_string(AZUL)
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        return p

    def add_texto_auto(text, **kwargs):
        """Escolhe automaticamente entre texto normal e link azul, conforme o conteúdo."""
        if is_path(text):
            return add_link_texto(text, **{k: v for k, v in kwargs.items() if k in ('size', 'space_after', 'space_before')})
        return add_texto(text, **kwargs)

    def add_bullet(text, size=11, space_after=7):
        """Item com marcador (•) — usado para nomes de script/arquivo."""
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.font.name = 'Arial Narrow'
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(space_after)
        return p

    # ─── CONSTRUÇÃO DE UMA PÁGINA/ETAPA ────────────────────────
    def criar_pagina_etapa(etapa_num, etapa_data, folha_num):

        # ── Bloco 1: cabeçalho principal (grid fino de 8 colunas) ──
        widths = [3.0, 2.5, 2.5, 4.0, 2.0, 2.0, 1.0, 1.0]  # soma = 18.0
        hdr = doc.add_table(rows=3, cols=8)
        hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(hdr, widths)
        r0, r1, r2 = hdr.rows[0], hdr.rows[1], hdr.rows[2]

        # Linha 1: PRODAM | CARACTERÍSTICA DE PROGRAMA/UTILITÁRIO | CÓDIGO DA ROTINA | ETAPA
        c_prodam = r0.cells[0]
        c_carac = r0.cells[1].merge(r0.cells[2]).merge(r0.cells[3])
        c_codrot = r0.cells[4].merge(r0.cells[5])
        c_etapa = r0.cells[6].merge(r0.cells[7])
        cell_text(c_prodam, 'PRODAM', 15, True, WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(c_carac, 'CARACTERÍSTICA DE PROGRAMA/UTILITÁRIO', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
        cell_label_value(c_codrot, 'CÓDIGO DA ROTINA', cod, 7, 12, WD_ALIGN_PARAGRAPH.CENTER)
        cell_label_value(c_etapa, 'ETAPA', str(etapa_num).zfill(2), 7, 12, WD_ALIGN_PARAGRAPH.CENTER)

        # Linha 2: SISTEMA | SUBSISTEMA
        c_sis = r1.cells[0].merge(r1.cells[1])
        c_sub = r1.cells[2].merge(r1.cells[3]).merge(r1.cells[4]).merge(r1.cells[5]).merge(r1.cells[6]).merge(r1.cells[7])
        cell_label_value(c_sis, 'SISTEMA', sis, 7, 10)
        cell_label_value(c_sub, 'SUBSISTEMA', sub, 7, 10)

        # Linha 3: CÓD. PROGRAMA | DENOMINAÇÃO DO PROGRAMA | DATA DA ELABORAÇÃO | FOLHA
        c_codprog = r2.cells[0]
        c_den = r2.cells[1].merge(r2.cells[2]).merge(r2.cells[3])
        c_data = r2.cells[4].merge(r2.cells[5])
        c_folha = r2.cells[6].merge(r2.cells[7])
        cell_label_value(c_codprog, 'CÓD. PROGRAMA', 'SQLPLUS', 7, 10)
        cell_label_value(c_den, 'DENOMINAÇÃO DO PROGRAMA', den, 7, 10)
        cell_label_value(c_data, 'DATA DA ELABORAÇÃO', dat, 7, 10)
        cell_label_value(c_folha, 'FOLHA', str(folha_num).zfill(2), 7, 10, WD_ALIGN_PARAGRAPH.CENTER)

        set_table_borders(hdr)

        # ── Bloco 2: TEMPO / MEMBER NAME / SORT INTERNO ──
        widths2 = [3.0, 3.0, 4.0, 4.0, 4.0]  # soma = 18.0
        tbl2 = doc.add_table(rows=2, cols=5)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(tbl2, widths2)
        a, b = tbl2.rows[0], tbl2.rows[1]

        c_tempo = a.cells[0].merge(b.cells[0])
        c_member = a.cells[1].merge(b.cells[1])
        c_sortint = a.cells[2].merge(a.cells[3]).merge(a.cells[4])

        c_tempo.text = ''
        p_t = c_tempo.paragraphs[0]
        p_t.paragraph_format.space_after = Pt(1)
        rt = p_t.add_run('TEMPO')
        rt.font.size = Pt(7); rt.font.name = 'Arial Narrow'; rt.font.color.rgb = RGBColor.from_string(PRETO)
        for lbl in ['☐ SORT', '☐ MERGE']:
            pc = c_tempo.add_paragraph()
            pc.paragraph_format.space_before = Pt(2)
            pc.paragraph_format.space_after = Pt(1)
            rc = pc.add_run(lbl)
            rc.font.size = Pt(11); rc.font.name = 'Arial Narrow'; rc.font.color.rgb = RGBColor.from_string(PRETO)
        c_tempo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell_label_value(c_member, 'MEMBER NAME', sub, 7, 10)
        cell_text(c_sortint, 'SORT INTERNO', 10, True, WD_ALIGN_PARAGRAPH.CENTER)

        for lbl, cell in [('TAM. REGISTRO:', b.cells[2]), ('NRO. REGISTRO:', b.cells[3]),
                          ('FUNÇÃO INCREMENTO:', b.cells[4])]:
            cell_text(cell, lbl, 8)

        set_table_borders(tbl2)

        # ── Bloco 3: barra FUNÇÕES DE CONTROLE ──
        tbl3 = doc.add_table(rows=1, cols=1)
        tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(tbl3, [18.0])
        cell_text(tbl3.rows[0].cells[0], 'FUNÇÕES DE CONTROLE (PARM, TESTE, RETURN CODE, UTILITÁRIOS)',
                  10, True, WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
        set_table_borders(tbl3)

        # ── Bloco 4: PARM = Campo1 + tabela de campos ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r1_ = p.add_run('PARM')
        r1_.font.underline = True
        r1_.font.size = Pt(11); r1_.font.name = 'Arial Narrow'; r1_.font.color.rgb = RGBColor.from_string(PRETO)
        r2_ = p.add_run(' = Campo1')
        r2_.font.size = Pt(11); r2_.font.name = 'Arial Narrow'; r2_.font.color.rgb = RGBColor.from_string(PRETO)

        widths4 = [2.4, 4.8, 3.6, 2.4, 4.8]  # soma = 18.0
        tbl4 = doc.add_table(rows=2, cols=5)
        tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(tbl4, widths4)
        headers4 = ['campo', 'Nome', 'Tamanho', 'Obrigat.', 'Observação']
        for i, h in enumerate(headers4):
            cell_text(tbl4.rows[0].cells[i], h, 8)
        for c in range(5):
            cell_text(tbl4.rows[1].cells[c], '', 8)
        set_table_borders(tbl4)

        # ── ÁREA DE CONTEÚDO (corpo da etapa) ──────────────────
        add_texto('', size=2, space_after=2)

        if etapa_num == 0:
            dia = dados.get('dia', '').strip()
            if freq == 'mensalmente' and dia:
                add_texto(f'Rotina que inicialmente será executada mensalmente dia {dia} às {hora} da manhã.')
            else:
                add_texto(f'Rotina que inicialmente será executada {freq}, às {hora} da manhã.')

            add_texto('OBJETIVO:', bold=True, underline=True, space_after=6)
            objetivo = dados.get('objetivo', '').strip()
            if objetivo:
                for line in objetivo.split('\n'):
                    line = line.strip()
                    if line:
                        add_texto_auto(line)
            else:
                add_texto('Automatizar a disponibilização do arquivo na pasta (CP):')
            # ETAPA 00 fica só com a rotina + objetivo — nada mais aqui.

        elif etapa_data.get('tipo') == 'SQLPLUS':
            add_texto('EXECUTAR SCRIPT SQL:', bold=True, underline=True, space_after=6)
            add_bullet(etapa_data.get('inp', ''))
            obs = etapa_data.get('obs', '')
            if obs:
                add_texto(f'Observação: {obs}')
            if etapa_data.get('out'):
                add_texto('ARQUIVO DE SAÍDA GERADO PELO SCRIPT SQL:', bold=True, underline=True,
                           space_before=14, space_after=6)
                add_bullet(etapa_data.get('out', ''))
                if '&' in etapa_data.get('out', ''):
                    var_fmt = 'MM-YYYY' if freq == 'mensalmente' else 'DD-MM-YYYY'
                    add_texto(f"Onde: &x = '{var_fmt}'")

            add_texto('COMANDOS NECESSÁRIOS ANTES DE INICIAR O SQLPLUS:', bold=True, underline=True,
                       space_before=14, space_after=6)
            for cmd in ['export LANG=pt_BR.ISO-8859-1',
                        'export NLS_LANG=BRAZILIAN PORTUGUESE_BRAZIL.WE8MSWIN1252']:
                add_texto(cmd, space_after=5)

        elif etapa_data.get('tipo') == 'COMPACTAR':
            add_texto('COMPACTAR ARQUIVO:', bold=True, underline=True, space_after=6)
            add_bullet(etapa_data.get('inp', ''))

        elif etapa_data.get('tipo') == 'TRANSFERIR':
            inp = etapa_data.get('inp', '')
            arquivos = inp if isinstance(inp, list) else ([inp] if inp else [])
            titulo = 'TRANSFERIR ARQUIVOS:' if len(arquivos) > 1 else 'TRANSFERIR ARQUIVO:'
            add_texto(titulo, bold=True, underline=True, space_after=6)
            for arq in arquivos:
                add_bullet(arq)

            destino_tipo = dados.get('destino_tipo', 'SFTP')
            ip_val = dados.get('ip', '').strip()
            if destino_tipo == 'CP' or not ip_val:
                caminho = (dados.get('cp_path', '').strip()
                           or dados.get('pasta', '').strip()
                           or extrair_path_do_texto(dados.get('objetivo', '')))
                add_texto('Disponibilizar os arquivos compactados, na pasta (CP):', space_before=14)
                add_link_texto(caminho)
            else:
                add_texto('Destino do arquivo gerado (servidor SFTP):', bold=True, underline=True,
                           space_before=14, space_after=6)
                sftp_items = [
                    f'- IP: {dados.get("ip", "")}',
                    f'- Porta: {dados.get("porta", "")}',
                    f'- Protocolo: {dados.get("proto", "SFTP")}',
                    f'- Pasta: {dados.get("pasta", "")}',
                    'Credenciais:',
                    f'- Usuário: {dados.get("user", "")}',
                    f'- Senha: {dados.get("senha", "")}',
                ]
                for item in sftp_items:
                    add_texto(item, space_after=5)

        # Quebra de página (exceto na última etapa)
        if etapa_num < len(etapas):
            doc.add_page_break()

    # ─── GERAR TODAS AS PÁGINAS ─────────────────────────────
    criar_pagina_etapa(0, {}, 1)
    for i, et in enumerate(etapas):
        criar_pagina_etapa(et['num'], et, i + 2)

    # ─── RODAPÉ FIXO (CAMPOS DE CLASSIFICAÇÃO + RELATÓRIOS) ─────
    # Usa o rodapé real do Word — assim repete automaticamente em
    # TODAS as páginas, sempre colado na base da folha.
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = ''

    NCOLS_CLASS = 12
    tbl_class = footer.add_table(rows=3, cols=NCOLS_CLASS, width=Cm(18))
    tbl_class.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci in range(1, NCOLS_CLASS):
        tbl_class.rows[0].cells[0].merge(tbl_class.rows[0].cells[ci])
    cell_text(tbl_class.rows[0].cells[0], 'CAMPOS DE CLASSIFICAÇÃO', 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    headers_class = ['NUM', 'NOME DO CAMPO', 'POS.REL.', 'DIM', 'FORM', 'SEQ.',
                      'NUM', 'NOME DO CAMPO', 'POS.REL.', 'DIM', 'FORM', 'SEQ.']
    for i, h in enumerate(headers_class):
        cell_text(tbl_class.rows[1].cells[i], h, 7)
    for c in range(NCOLS_CLASS):
        cell_text(tbl_class.rows[2].cells[c], '', 7)
    set_table_borders(tbl_class)

    p_gap = footer.add_paragraph()
    p_gap.paragraph_format.space_after = Pt(2)
    p_gap.paragraph_format.space_before = Pt(2)

    NCOLS_REL = 9
    tbl_rel = footer.add_table(rows=3, cols=NCOLS_REL, width=Cm(18))
    tbl_rel.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci in range(1, NCOLS_REL):
        tbl_rel.rows[0].cells[0].merge(tbl_rel.rows[0].cells[ci])
    cell_text(tbl_rel.rows[0].cells[0], 'RELATÓRIOS', 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    headers_rel = ['DSNAME/ETAPA', 'CS', 'SQ', 'CAMPO DE CONTROLE', 'CADEIA ESP.',
                   'CÓD. FORMULÁRIO', 'CÓPIAS', 'IMPRESSORA', 'FCB']
    for i, h in enumerate(headers_rel):
        cell_text(tbl_rel.rows[1].cells[i], h, 7)
    for c in range(NCOLS_REL):
        cell_text(tbl_rel.rows[2].cells[c], '', 7)
    set_table_borders(tbl_rel)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ═══════════════════════════════════════════════════════════════════
# GERADOR DE SVG (FLUXOGRAMA) — mantém lógica existente
# ═══════════════════════════════════════════════════════════════════

def escape_svg(s):
    return str(s or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def wrap_text(text, max_chars):
    words = str(text or '').split(' ')
    lines, cur = [], ''
    for w in words:
        if len((cur + ' ' + w).strip()) > max_chars:
            if cur:
                lines.append(cur)
            cur = w
        else:
            cur = (cur + ' ' + w).strip()
    if cur:
        lines.append(cur)
    return lines


def gerar_svg(dados: dict) -> list:
    """Gera lista de SVGs (um por folha) com paginação automática."""
    cod = dados.get('cod', '???')
    sis = dados.get('sis', '')
    sub = dados.get('sub', '')
    den = dados.get('den', '')
    amb = dados.get('amb', 'UNIX')
    dat = dados.get('dat', '')
    hora = dados.get('hora', '??h??')
    ip = dados.get('ip', '')
    porta = dados.get('porta', '')
    pasta = dados.get('pasta', '')
    user = dados.get('user', '')
    destino_tipo = dados.get('destino_tipo', 'SFTP')
    cp_path = (dados.get('cp_path', '') or dados.get('pasta', '') or extrair_path_do_texto(dados.get('objetivo', ''))).strip()
    etapas = dados.get('etapas', [])
    tabelas = dados.get('tabelas', [])

    W = 700
    BLU, GRY, LGY, BRD = '#1F4E79', '#333', '#888', '#333'
    AZUL = '#0563C1'
    FNT = 'Arial Narrow,Arial,sans-serif'
    ex = escape_svg

    tem_transferir = any(e.get('tipo') == 'TRANSFERIR' for e in etapas)
    tem_sftp = tem_transferir and destino_tipo == 'SFTP' and bool(ip.strip())
    tem_cp = tem_transferir and not tem_sftp and bool(cp_path)

    # ─── Estimativa de altura para decidir paginação ──────────
    # Altura max do SVG que cabe numa A4 no Word (27.7cm úteis × 700/18)
    MAX_SVG_H = 1077
    HEADER_H  = 116   # header PRODAM (96px) + gap (20px)
    FOOTER_H  = 34    # linha + texto + margem
    FIM_H     = 74    # seta + pill FIM
    SFTP_H    = 162   # seta + cilindro SFTP
    CP_H      = 70    # seta + oval com caminho azul
    TABELAS_H = 16    # linha de tabelas Oracle
    # Por etapa: seta(28) + oval_in(~58) + box(56) + oval_out(~58) + gap(10)
    ET_H      = 210   # estimativa conservadora

    n_et = len(etapas)

    # Estima altura total se tudo ficasse em 1 folha
    est_total = HEADER_H + (ET_H * n_et) + FIM_H + FOOTER_H + TABELAS_H
    if tem_sftp:
        est_total += SFTP_H
    elif tem_cp:
        est_total += CP_H

    # Se cabe em 1 página A4 → tudo numa folha; senão → max 2 por folha
    MAX_ET = n_et if (n_et > 0 and est_total <= MAX_SVG_H) else 2

    # Agrupa etapas em folhas
    folhas_etapas = []
    for i in range(0, len(etapas), MAX_ET):
        folhas_etapas.append(etapas[i:i + MAX_ET])
    if not folhas_etapas:
        folhas_etapas = [[]]

    # Se tudo cabe em 1 folha (MAX_ET == n_et), a folha NÃO está "cheia"
    # para fins de paginação — o destino (SFTP/CP) deve ficar na mesma folha
    ultima_cheia = len(folhas_etapas[-1]) >= MAX_ET and MAX_ET != n_et
    tem_destino = tem_sftp or tem_cp
    total_folhas = len(folhas_etapas) + (1 if tem_destino and ultima_cheia else 0)

    def draw_header(els, Y, fnum):
        PRETO = '#000'
        C1, C2, C3 = 130, 390, 180  # soma = 700
        RH1, RH2, RH3 = 36, 26, 40
        HH = RH1 + RH2 + RH3
        bw = 1.3
        els.append(f'<rect x="0" y="{Y}" width="{W}" height="{HH}" fill="#fff" stroke="{PRETO}" stroke-width="{bw}"/>')

        # Linha 1: PRODAM | FLUXOGRAMA DE ROTINA DE OPERAÇÃO | CÓDIGO DA ROTINA
        els.append(f'<line x1="{C1}" y1="{Y}" x2="{C1}" y2="{Y+RH1}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<line x1="{C1+C2}" y1="{Y}" x2="{C1+C2}" y2="{Y+RH1}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<line x1="0" y1="{Y+RH1}" x2="{W}" y2="{Y+RH1}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<text x="{C1//2}" y="{Y+RH1//2+6}" font-size="18" font-weight="bold" text-anchor="middle" fill="{PRETO}" font-family="{FNT}">PRODAM</text>')
        els.append(f'<text x="{C1+C2//2}" y="{Y+RH1//2+6}" font-size="15" font-weight="bold" text-anchor="middle" fill="{PRETO}" font-family="{FNT}">FLUXOGRAMA DE ROTINA DE OPERAÇÃO</text>')
        els.append(f'<text x="{C1+C2+C3//2}" y="{Y+14}" font-size="10" text-anchor="middle" fill="{PRETO}" font-family="{FNT}">CÓDIGO DA ROTINA:</text>')
        els.append(f'<text x="{C1+C2+C3//2}" y="{Y+29}" font-size="14" font-weight="bold" text-anchor="middle" fill="{PRETO}" font-family="{FNT}">{ex(cod)} ({ex(amb)})</text>')

        # Linha 2: SISTEMA | SUBSISTEMA
        y2 = Y + RH1
        SW1 = 400
        els.append(f'<line x1="{SW1}" y1="{y2}" x2="{SW1}" y2="{y2+RH2}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<line x1="0" y1="{y2+RH2}" x2="{W}" y2="{y2+RH2}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<text x="8" y="{y2+11}" font-size="10" fill="{PRETO}" font-family="{FNT}">SISTEMA:</text>')
        els.append(f'<text x="8" y="{y2+RH2-4}" font-size="13" font-weight="bold" fill="{PRETO}" font-family="{FNT}">{ex(sis)}</text>')
        els.append(f'<text x="{SW1+8}" y="{y2+11}" font-size="10" fill="{PRETO}" font-family="{FNT}">SUBSISTEMA:</text>')
        els.append(f'<text x="{SW1+8}" y="{y2+RH2-4}" font-size="13" font-weight="bold" fill="{PRETO}" font-family="{FNT}">{ex(sub)}</text>')

        # Linha 3: DENOMINAÇÃO DO PROGRAMA | DATA DA ELABORAÇÃO | FOLHA
        y3 = y2 + RH2
        DW = 430
        FW3 = 140
        els.append(f'<line x1="{DW}" y1="{y3}" x2="{DW}" y2="{y3+RH3}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<line x1="{DW+FW3}" y1="{y3}" x2="{DW+FW3}" y2="{y3+RH3}" stroke="{PRETO}" stroke-width="{bw}"/>')
        els.append(f'<text x="8" y="{y3+12}" font-size="10" fill="{PRETO}" font-family="{FNT}">DENOMINAÇÃO DO PROGRAMA</text>')
        els.append(f'<text x="8" y="{y3+28}" font-size="13" font-weight="bold" fill="{PRETO}" font-family="{FNT}">{ex(den)}</text>')
        els.append(f'<text x="{DW+8}" y="{y3+12}" font-size="10" fill="{PRETO}" font-family="{FNT}">DATA DA ELABORAÇÃO:</text>')
        els.append(f'<text x="{DW+8}" y="{y3+28}" font-size="12" fill="{PRETO}" font-family="{FNT}">{ex(dat)}</text>')
        els.append(f'<text x="{DW+FW3+8}" y="{y3+12}" font-size="10" fill="{PRETO}" font-family="{FNT}">FOLHA:</text>')
        els.append(f'<text x="{DW+FW3+8}" y="{y3+28}" font-size="14" font-weight="bold" fill="{PRETO}" font-family="{FNT}">{str(fnum).zfill(2)}</text>')

        return Y + HH + 20

    def draw_oval(els, Y, nome):
        # Aceita lista (vários arquivos) ou string — quebra em várias linhas se precisar
        if isinstance(nome, (list, tuple)):
            linhas = list(nome)
        else:
            linhas = wrap_text(str(nome or ''), 55)
        if not linhas:
            linhas = ['']
        maior = max(len(l) for l in linhas)
        fw = min(W - 60, max(140, maior * 6 + 50))
        line_h = 14
        fh = max(30, 16 + line_h * len(linhas))
        cy = Y + fh // 2
        els.append(f'<ellipse cx="{W//2}" cy="{cy}" rx="{fw//2}" ry="{fh//2}" fill="none" stroke="{BRD}" stroke-width="1.2"/>')
        start_y = cy - (line_h * (len(linhas) - 1)) // 2 + 4
        for i, linha in enumerate(linhas):
            els.append(f'<text x="{W//2}" y="{start_y + i*line_h}" font-size="12" text-anchor="middle" fill="{GRY}" font-family="{FNT}">{ex(linha)}</text>')
        return Y + fh

    def draw_flag(els, Y, texto, tabelas_lista=None):
        """Forma de 'fita/bandeirola' (ponta arredondada à esquerda, recorte em V
        à direita) — usada para reexibir o script/arquivo junto da lista de tabelas."""
        linhas = wrap_text(str(texto or ''), 55) or ['']
        maior = max(len(l) for l in linhas)
        fw = min(W - 160, max(280, maior * 6 + 60))
        line_h = 14
        fh = max(40, 20 + line_h * len(linhas))
        X = (W - fw) // 2 - 40
        topY, botY = Y, Y + fh
        midY = Y + fh / 2
        right = X + fw
        notch = right - fw * 0.09
        els.append(
            f'<path d="M {X},{midY} '
            f'Q {X+fw*0.05},{topY} {X+fw*0.22},{topY} '
            f'L {right},{topY} '
            f'L {notch},{midY} '
            f'L {right},{botY} '
            f'L {X+fw*0.22},{botY} '
            f'Q {X+fw*0.05},{botY} {X},{midY} Z" '
            f'fill="#fff" stroke="{BRD}" stroke-width="1.2"/>'
        )
        start_y = midY - (line_h * (len(linhas) - 1)) / 2 + 4
        for i, linha in enumerate(linhas):
            els.append(f'<text x="{X+fw*0.46}" y="{start_y + i*line_h}" font-size="12" text-anchor="middle" fill="{GRY}" font-family="{FNT}">{ex(linha)}</text>')

        bottom = Y + fh
        if tabelas_lista:
            tx = right + 14
            ty = Y + 10
            els.append(f'<text x="{tx}" y="{ty}" font-size="10" fill="{LGY}" font-family="{FNT}">Tabelas:</text>')
            for ti, tn in enumerate(tabelas_lista):
                els.append(f'<text x="{tx}" y="{ty + 11 + ti*11}" font-size="10" fill="{GRY}" font-family="{FNT}">{ex(tn)}</text>')
            bottom = max(bottom, ty + 11 + 11 * len(tabelas_lista))
        return Y + fh, bottom

    def draw_arrow(els, Y, length=24):
        mid = Y + length - 8
        els.append(f'<line x1="{W//2}" y1="{Y}" x2="{W//2}" y2="{mid}" stroke="{GRY}" stroke-width="1.2"/>')
        els.append(f'<polygon points="{W//2},{Y+length} {W//2-5},{mid} {W//2+5},{mid}" fill="{GRY}"/>')
        return Y + length

    def draw_footer(els, Y, fnum):
        els.append(f'<line x1="0" y1="{Y}" x2="{W}" y2="{Y}" stroke="{LGY}" stroke-width="0.5"/>')
        els.append(f'<text x="10" y="{Y+12}" font-size="10" fill="{LGY}" font-family="{FNT}">PRODAM / {ex(sis)} / {ex(cod)}</text>')
        els.append(f'<text x="{W-10}" y="{Y+12}" font-size="10" text-anchor="end" fill="{LGY}" font-family="{FNT}">FOLHA {str(fnum).zfill(2)}/{str(total_folhas).zfill(2)}</text>')
        return Y + 20

    def draw_sftp(els, Y):
        CW_s, CRH_s = 260, 120
        CX_s = (W - CW_s) // 2
        CR = 16
        els.append(f'<ellipse cx="{CX_s+CW_s//2}" cy="{Y+CR}" rx="{CW_s//2}" ry="{CR}" fill="none" stroke="{BRD}" stroke-width="1.5"/>')
        els.append(f'<rect x="{CX_s}" y="{Y+CR}" width="{CW_s}" height="{CRH_s-CR*2}" fill="#fff" stroke="{BRD}" stroke-width="1.5"/>')
        els.append(f'<ellipse cx="{CX_s+CW_s//2}" cy="{Y+CRH_s-CR}" rx="{CW_s//2}" ry="{CR}" fill="none" stroke="{BRD}" stroke-width="1.5"/>')
        els.append(f'<text x="{CX_s+CW_s//2}" y="{Y+38}" font-size="14" font-weight="bold" text-anchor="middle" fill="{GRY}" font-family="{FNT}">SFTP — Servidor Externo</text>')
        els.append(f'<text x="{CX_s+CW_s//2}" y="{Y+55}" font-size="12" text-anchor="middle" fill="{GRY}" font-family="{FNT}">IP: {ex(ip)}  |  Porta: {ex(porta)}</text>')
        els.append(f'<text x="{CX_s+CW_s//2}" y="{Y+70}" font-size="12" text-anchor="middle" fill="{GRY}" font-family="{FNT}">Usuário: {ex(user)}</text>')
        els.append(f'<text x="{CX_s+CW_s//2}" y="{Y+85}" font-size="11" text-anchor="middle" fill="{LGY}" font-family="{FNT}">{ex(pasta)}</text>')
        return Y + CRH_s + 10

    def draw_cp_path(els, Y, path):
        """Caixa com o caminho da pasta de rede (CP), em azul — destino sem SFTP."""
        linhas = wrap_text(str(path or ''), 60)
        if not linhas:
            linhas = ['']
        line_h = 14
        fh = max(34, 20 + line_h * len(linhas))
        CW_c = min(W - 60, max(220, max(len(l) for l in linhas) * 6 + 40))
        CX_c = (W - CW_c) // 2
        els.append(f'<rect x="{CX_c}" y="{Y}" width="{CW_c}" height="{fh}" rx="8" fill="none" stroke="{BRD}" stroke-width="1.2"/>')
        els.append(f'<text x="{W//2}" y="{Y+14}" font-size="11" text-anchor="middle" fill="{LGY}" font-family="{FNT}">Disponibilizar na pasta (CP):</text>')
        start_y = Y + 28
        for i, linha in enumerate(linhas):
            els.append(f'<text x="{W//2}" y="{start_y + i*line_h}" font-size="11" text-anchor="middle" fill="{AZUL}" text-decoration="underline" font-family="{FNT}">{ex(linha)}</text>')
        return Y + fh + 10

    def draw_fim(els, Y):
        Y = draw_arrow(els, Y)
        Y += 4
        FW, FX = 160, (W - 160) // 2
        els.append(f'<rect x="{FX}" y="{Y}" width="{FW}" height="32" rx="16" fill="{BLU}"/>')
        els.append(f'<text x="{W//2}" y="{Y+21}" font-size="14" font-weight="bold" text-anchor="middle" fill="#fff" font-family="{FNT}">FIM</text>')
        return Y + 46

    def make_svg(els, Y):
        return f'<svg viewBox="0 0 {W} {Y}" xmlns="http://www.w3.org/2000/svg" width="{W}" style="background:#fff;max-width:100%">\n' + '\n'.join(els) + '\n</svg>'

    # ─── GERA FOLHAS ─────────────────────────────────────────────
    svgs = []
    prev_out = None

    for fi, folha_ets in enumerate(folhas_etapas):
        fnum = fi + 1
        els = []
        Y = draw_header(els, 0, fnum)
        folha_extra_bottom = 0  # maior "vazamento" à direita (lista de tabelas) nesta folha

        for et in folha_ets:
            EW, EX = 200, (W - 200) // 2
            EH = 56

            Y = draw_arrow(els, Y)
            Y += 4

            inp = et.get('inp', '')
            if inp and inp != prev_out:
                Y = draw_oval(els, Y, inp)
                Y += 4
                Y = draw_arrow(els, Y, 20)
                Y += 4

            if et['tipo'] == 'SQLPLUS':
                RW, RH = 90, 48
                RX = EX - RW - 30
                RY = Y + (EH - RH) // 2
                els.append(f'<rect x="{RX}" y="{RY}" width="{RW}" height="{RH}" rx="10" fill="none" stroke="{BRD}" stroke-width="1.2"/>')
                els.append(f'<text x="{RX+RW//2}" y="{RY+18}" font-size="13" font-weight="bold" text-anchor="middle" fill="{GRY}" font-family="{FNT}">Rotina</text>')
                els.append(f'<text x="{RX+RW//2}" y="{RY+34}" font-size="13" font-weight="bold" text-anchor="middle" fill="{BLU}" font-family="{FNT}">{ex(cod)}</text>')
                els.append(f'<line x1="{RX+RW}" y1="{RY+RH//2}" x2="{EX}" y2="{Y+EH//2}" stroke="{BRD}" stroke-width="1.2"/>')
                els.append(f'<polygon points="{EX},{Y+EH//2} {EX-7},{Y+EH//2-4} {EX-7},{Y+EH//2+4}" fill="{BRD}"/>')
                CW, CRH = 90, 56
                CX = EX + EW + 30
                CY = Y + (EH - CRH) // 2
                CR = 10
                els.append(f'<ellipse cx="{CX+CW//2}" cy="{CY+CR}" rx="{CW//2}" ry="{CR}" fill="none" stroke="{BRD}" stroke-width="1.2"/>')
                els.append(f'<rect x="{CX}" y="{CY+CR}" width="{CW}" height="{CRH-CR*2}" fill="#fff" stroke="{BRD}" stroke-width="1.2"/>')
                els.append(f'<ellipse cx="{CX+CW//2}" cy="{CY+CRH-CR}" rx="{CW//2}" ry="{CR}" fill="none" stroke="{BRD}" stroke-width="1.2"/>')
                els.append(f'<text x="{CX+CW//2}" y="{CY+CRH//2}" font-size="12" text-anchor="middle" fill="{GRY}" font-family="{FNT}">Tabelas</text>')
                els.append(f'<text x="{CX+CW//2}" y="{CY+CRH//2+14}" font-size="12" text-anchor="middle" fill="{GRY}" font-family="{FNT}">Oracle</text>')
                els.append(f'<line x1="{CX}" y1="{CY+CRH//2}" x2="{EX+EW}" y2="{Y+EH//2}" stroke="{BRD}" stroke-width="1.2"/>')
                els.append(f'<polygon points="{EX+EW},{Y+EH//2} {EX+EW+7},{Y+EH//2-4} {EX+EW+7},{Y+EH//2+4}" fill="{BRD}"/>')

            els.append(f'<rect x="{EX}" y="{Y}" width="{EW}" height="{EH}" fill="none" stroke="{BRD}" stroke-width="1.8"/>')
            els.append(f'<line x1="{EX}" y1="{Y+28}" x2="{EX+EW}" y2="{Y+28}" stroke="{BRD}" stroke-width="1"/>')
            els.append(f'<text x="{EX+EW//2}" y="{Y+20}" font-size="15" font-weight="bold" text-anchor="middle" fill="{GRY}" font-family="{FNT}">ETAPA {str(et["num"]).zfill(2)}</text>')
            els.append(f'<text x="{EX+EW//2}" y="{Y+47}" font-size="13" text-anchor="middle" fill="{GRY}" font-family="{FNT}">{ex(et.get("titulo","") or et["tipo"])}</text>')
            Y += EH

            # Reexibe o script + lista de tabelas na forma de "fita", como no modelo PRODAM
            if et['tipo'] == 'SQLPLUS':
                Y += 4
                Y = draw_arrow(els, Y, 20)
                Y += 4
                Y, flag_bottom = draw_flag(els, Y, et.get('inp', ''), et.get('tabelas', []))
                folha_extra_bottom = max(folha_extra_bottom, flag_bottom)

            out = et.get('out', '')
            if out:
                Y += 4
                Y = draw_arrow(els, Y, 20)
                Y += 4
                Y = draw_oval(els, Y, out)

            prev_out = out
            Y += 10

        # Se alguma lista de tabelas "vazou" pra baixo do fluxo normal desta
        # folha, garante que o restante (destino/FIM/rodapé) comece depois dela.
        Y = max(Y, folha_extra_bottom + 10)

        # Destino (SFTP ou pasta CP) nesta folha, se couber
        is_last = fi == len(folhas_etapas) - 1
        if tem_destino and is_last and not ultima_cheia:
            Y = draw_arrow(els, Y, 28)
            Y += 4
            if tem_sftp:
                Y = draw_sftp(els, Y)
            else:
                Y = draw_cp_path(els, Y, cp_path)
            Y = draw_fim(els, Y)
        elif is_last and not tem_destino:
            Y = draw_fim(els, Y)

        Y += 10
        Y = draw_footer(els, Y, fnum)
        svgs.append(make_svg(els, Y + 4))

    # Folha extra para o destino (SFTP/CP) se a última folha de etapas estava cheia
    if tem_destino and ultima_cheia:
        els = []
        Y = draw_header(els, 0, total_folhas)
        if prev_out:
            Y = draw_arrow(els, Y)
            Y += 4
            Y = draw_oval(els, Y, prev_out)
            Y += 4
        Y = draw_arrow(els, Y, 28)
        Y += 4
        if tem_sftp:
            Y = draw_sftp(els, Y)
        else:
            Y = draw_cp_path(els, Y, cp_path)
        Y = draw_fim(els, Y)
        Y += 10
        Y = draw_footer(els, Y, total_folhas)
        svgs.append(make_svg(els, Y + 4))

    return svgs


# ═══════════════════════════════════════════════════════════════════
# GERADOR DE FLUXOGRAMA EM WORD (.docx)
# ═══════════════════════════════════════════════════════════════════

def gerar_fluxograma_docx(svgs: list, cod: str = 'fluxograma') -> bytes:
    """Converte lista de SVGs do fluxograma em .docx.
    Se o conteudo total couber em uma pagina A4, combina em uma imagem so.
    Caso contrario, usa uma pagina Word por folha SVG.
    """
    import cairosvg
    from docx import Document
    from docx.shared import Cm

    PAGE_W_CM = 18.0   # largura util: 21 - 1.5 - 1.5
    PAGE_H_CM = 27.7   # altura util:  29.7 - 1 - 1
    SVG_W_PX  = 700    # viewBox width padrao dos SVGs gerados
    GAP_PX    = 30     # espaco entre folhas no SVG combinado

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    def svg_height(svg_str):
        m = re.search(r'viewBox="0 0 \d+ (\d+)"', svg_str)
        return int(m.group(1)) if m else 500

    def combine_svgs(svg_list):
        heights = [svg_height(s) for s in svg_list]
        total_h = sum(heights) + (len(svg_list) - 1) * GAP_PX
        els = []
        offset = 0
        for s, h in zip(svg_list, heights):
            inner = re.search(r'<svg[^>]*>(.*)</svg>', s, re.DOTALL)
            if inner:
                els.append(f'<g transform="translate(0,{offset})">{inner.group(1)}</g>')
            offset += h + GAP_PX
        return (
            f'<svg viewBox="0 0 {SVG_W_PX} {total_h}" '
            f'xmlns="http://www.w3.org/2000/svg" width="{SVG_W_PX}" style="background:#fff">\n'
            + '\n'.join(els) + '\n</svg>'
        )

    if len(svgs) == 1:
        png_bytes = cairosvg.svg2png(bytestring=svgs[0].encode('utf-8'), output_width=1400)
        doc.add_picture(io.BytesIO(png_bytes), width=Cm(PAGE_W_CM))
    else:
        heights = [svg_height(s) for s in svgs]
        total_h_px = sum(heights) + (len(svgs) - 1) * GAP_PX
        combined_h_cm = total_h_px * (PAGE_W_CM / SVG_W_PX)

        if combined_h_cm <= PAGE_H_CM:
            # Cabe numa pagina: uma unica imagem combinada
            combined_svg = combine_svgs(svgs)
            png_bytes = cairosvg.svg2png(
                bytestring=combined_svg.encode('utf-8'), output_width=1400
            )
            doc.add_picture(io.BytesIO(png_bytes), width=Cm(PAGE_W_CM))
        else:
            # Nao cabe: uma pagina Word por folha SVG
            for i, svg_str in enumerate(svgs):
                png_bytes = cairosvg.svg2png(
                    bytestring=svg_str.encode('utf-8'), output_width=1400
                )
                doc.add_picture(io.BytesIO(png_bytes), width=Cm(PAGE_W_CM))
                if i < len(svgs) - 1:
                    doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════
# MONTAGEM DAS ETAPAS A PARTIR DO SQL
# ═══════════════════════════════════════════════════════════════════

def montar_etapas(lista_dados_sql, compactar: bool, transferir: bool) -> list:
    """Monta a lista de etapas combinando UM OU MAIS scripts SQL numa única
    sequência (ex: script1→SQLPLUS→COMPACTAR, script2→SQLPLUS→COMPACTAR, ...,
    seguido de uma única etapa TRANSFERIR ao final com todos os .zip gerados).

    Aceita tanto uma lista de dicts (vários scripts) quanto um único dict
    (compatibilidade retroativa)."""
    if isinstance(lista_dados_sql, dict):
        lista_dados_sql = [lista_dados_sql]

    etapas = []
    zips_gerados = []
    num = 1

    for dados_sql in lista_dados_sql:
        arquivo_saida = dados_sql.get('arquivo_saida', '')
        obs_sql = descrever_operacao_sql(dados_sql)
        etapas.append({
            'num': num,
            'tipo': 'SQLPLUS',
            'titulo': 'Executar Script SQL',
            'inp': dados_sql.get('nome_script', ''),
            'out': arquivo_saida,
            'obs': obs_sql,
            'tabelas': dados_sql.get('tabelas', []),
        })
        num += 1

        if compactar:
            out_compact = arquivo_saida + '.zip' if arquivo_saida else ''
            etapas.append({
                'num': num,
                'tipo': 'COMPACTAR',
                'titulo': 'Compactar arquivo',
                'inp': arquivo_saida,
                'out': out_compact,
                'obs': '',
            })
            if out_compact:
                zips_gerados.append(out_compact)
            num += 1

    # Etapa final única — TRANSFERIR (todos os arquivos compactados de uma vez)
    if transferir and zips_gerados:
        etapas.append({
            'num': num,
            'tipo': 'TRANSFERIR',
            'titulo': 'Transferir arquivo(s)',
            'inp': zips_gerados,
            'out': '',
            'obs': '',
        })

    return etapas


# ═══════════════════════════════════════════════════════════════════
# INTERFACE STREAMLIT
# ═══════════════════════════════════════════════════════════════════

# Estado
if 'resultados' not in st.session_state:
    st.session_state.resultados = []

col_form, col_output = st.columns([1, 1.5])

with col_form:
    st.markdown("### 📂 Upload do(s) Script(s) SQL")

    st.markdown(
        '<div class="info-box">Faça upload de <strong>um ou mais scripts SQL</strong> (.sql) que fazem parte '
        'da <strong>mesma rotina</strong>. Todos entram como etapas sequenciais de um único documento '
        '(ex: 2 scripts → ETAPA 00 + 01/02 [script 1] + 03/04 [script 2] + 05 TRANSFERIR).</div>',
        unsafe_allow_html=True
    )

    sql_files = st.file_uploader(
        "Scripts SQL (.sql / .txt) — serão ordenados pelo nome do arquivo",
        type=['sql', 'txt'],
        key='sql_upload',
        accept_multiple_files=True,
    )

    st.markdown("---")
    st.markdown("### ⚙ Dados da Rotina")

    c1, c2 = st.columns(2)
    with c1:
        cod_rotina = st.text_input("Código da Rotina", value="", placeholder="Ex: SH07680Z")
    with c2:
        amb = st.selectbox("Ambiente", ['UNIX', 'Windows'], index=0)

    c1, c2 = st.columns(2)
    with c1:
        sistema = st.text_input("Sistema", value="SIGA SAUDE")
    with c2:
        subsistema = st.text_input("Subsistema", value="", placeholder="Ex: SH0768 / SIGA SAÚDE")

    denominacao = st.text_input("Denominação do Programa", value="", placeholder="Ex: Automatização - Relatório...")

    c1, c2 = st.columns(2)
    with c1:
        # Mês em português sem depender de locale
        import datetime as _dt
        _meses = {1:'Janeiro',2:'Fevereiro',3:'Março',4:'Abril',5:'Maio',6:'Junho',
                  7:'Julho',8:'Agosto',9:'Setembro',10:'Outubro',11:'Novembro',12:'Dezembro'}
        _hoje = _dt.date.today()
        _data_default = f"{_meses[_hoje.month]} {_hoje.year}"
        data_elab = st.text_input("Data da Elaboração", value=_data_default, placeholder="Ex: Junho 2026")
    with c2:
        horario = st.text_input("Horário de Execução", value="00h00", placeholder="Ex: 00h00")

    frequencia = st.selectbox(
        "Frequência de execução",
        ['diariamente', 'semanalmente', 'mensalmente', 'sob demanda'],
        index=0,
    )
    dia_mes = ''
    if frequencia == 'mensalmente':
        dia_mes = st.text_input("Dia do mês", value="", placeholder="Ex: 10")

    objetivo_txt = st.text_area(
        "Objetivo (texto livre)",
        value="Automatizar a disponibilização do arquivo na pasta (CP):",
        height=100,
        help="Vai direto no corpo da Característica, na ETAPA 00 — que fica só com a rotina + objetivo. "
             "Uma linha por parágrafo. Linhas que começam com '/' (caminho de rede) saem em azul, como link.",
    )

    st.markdown("---")
    st.markdown("### 📦 Etapas adicionais")

    incluir_compactar = st.checkbox("Incluir etapa COMPACTAR (ZIP) após cada script", value=True)
    incluir_transferir = st.checkbox("Incluir etapa TRANSFERIR (uma única, ao final)", value=True)

    destino_tipo = 'SFTP'
    sftp_ip = sftp_porta = sftp_pasta = sftp_user = sftp_senha = ''
    sftp_proto = 'SFTP'
    cp_path = ''

    if incluir_transferir:
        destino_tipo = st.radio(
            "Tipo de destino",
            ['SFTP', 'CP'],
            format_func=lambda x: 'Servidor externo (SFTP/FTP)' if x == 'SFTP' else 'Pasta de rede interna (CP)',
            horizontal=True,
        )
        if destino_tipo == 'SFTP':
            st.markdown("#### 🔗 Dados SFTP")
            c1, c2 = st.columns(2)
            with c1:
                sftp_ip = st.text_input("IP do servidor", placeholder="Ex: 144.22.186.214")
                sftp_proto = st.selectbox("Protocolo", ['SFTP', 'FTP'], index=0)
                sftp_user = st.text_input("Usuário", placeholder="Ex: prodam")
            with c2:
                sftp_porta = st.text_input("Porta", placeholder="Ex: 44322")
                sftp_pasta = st.text_input("Pasta destino", placeholder="Ex: /dados/relatorios")
                sftp_senha = st.text_input("Senha", type="password", placeholder="")
        else:
            _obj_path_guess = ''
            for _line in objetivo_txt.split('\n'):
                _line = _line.strip()
                if _line.startswith('/'):
                    _obj_path_guess = _line
                    break
            cp_path = st.text_input(
                "Caminho da pasta (CP)", value=_obj_path_guess,
                placeholder="Ex: /cifs/sh0743/SIGA_MENSAIS/REL_187_...",
                help="Pré-preenchido com o caminho já digitado no Objetivo (se houver). Pode editar.",
            )

    st.markdown("---")

    if st.button("▶▶ Gerar Característica + Fluxograma", type="primary", use_container_width=True):
        if not sql_files:
            st.error("Faça upload de pelo menos um script SQL.")
        elif not cod_rotina.strip():
            st.error("Preencha o Código da Rotina.")
        else:
            # Ordenar scripts por nome para manter a sequência (01_, 02_, ...)
            sql_files_sorted = sorted(sql_files, key=lambda f: f.name)

            lista_dados_sql = []
            for sql_file in sql_files_sorted:
                texto_sql = sql_file.read().decode('utf-8', errors='replace')
                lista_dados_sql.append(extrair_dados_sql(texto_sql, sql_file.name))

            cod_atual = cod_rotina.strip()

            # Monta as etapas combinando TODOS os scripts numa sequência única
            etapas = montar_etapas(lista_dados_sql, incluir_compactar, incluir_transferir)

            # União das tabelas de todos os scripts (para o resumo)
            tabelas_todas = sorted({t for ds in lista_dados_sql for t in ds['tabelas']})

            dados = {
                'cod': cod_atual,
                'sis': sistema,
                'sub': subsistema,
                'den': denominacao or f'Rotina {cod_atual}',
                'amb': amb,
                'dat': data_elab,
                'hora': horario,
                'freq': frequencia,
                'dia': dia_mes,
                'objetivo': objetivo_txt,
                'destino_tipo': destino_tipo,
                'cp_path': cp_path,
                'ip': sftp_ip if (incluir_transferir and destino_tipo == 'SFTP') else '',
                'porta': sftp_porta if (incluir_transferir and destino_tipo == 'SFTP') else '',
                'pasta': sftp_pasta if (incluir_transferir and destino_tipo == 'SFTP') else '',
                'user': sftp_user if (incluir_transferir and destino_tipo == 'SFTP') else '',
                'senha': sftp_senha if (incluir_transferir and destino_tipo == 'SFTP') else '',
                'proto': sftp_proto if (incluir_transferir and destino_tipo == 'SFTP') else 'SFTP',
                'etapas': etapas,
                'tabelas': tabelas_todas,
            }

            # Gerar Característica .docx
            docx_bytes = gerar_caracteristica_docx(dados)

            # Gerar Fluxograma SVG
            svgs = gerar_svg(dados)

            # Gerar Fluxograma .docx (SVG → PNG → Word)
            try:
                fluxo_docx_bytes = gerar_fluxograma_docx(svgs, cod_atual)
            except Exception as e:
                fluxo_docx_bytes = None
                st.warning(f"Fluxograma .docx não gerado (cairosvg indisponível): {e}")

            resultado = {
                'cod': cod_atual,
                'nome_script': ', '.join(ds['nome_script'] for ds in lista_dados_sql),
                'dados': dados,
                'dados_sql_list': lista_dados_sql,
                'docx_bytes': docx_bytes,
                'fluxo_docx_bytes': fluxo_docx_bytes,
                'svgs': svgs,
            }
            st.session_state.resultados = [resultado]

            st.markdown(
                f'<div class="success-box">✅ <strong>{len(lista_dados_sql)} script(s)</strong> combinados — '
                f'{len(tabelas_todas)} tabelas, '
                f'{len(etapas)} etapas (00 a {str(len(etapas)).zfill(2)})</div>',
                unsafe_allow_html=True
            )

    # Mostra dados extraídos
    if st.session_state.resultados:
        res = st.session_state.resultados[0]
        if not isinstance(res, dict) or 'dados_sql_list' not in res or 'cod' not in res:
            # Estado de uma versão antiga do app ficou salvo na sessão — descarta.
            st.session_state.resultados = []
        else:
            d = res.get('dados', {})
            with st.expander(f"📋 {res.get('cod', '?')} — {len(res.get('dados_sql_list', []))} script(s)", expanded=False):
                for ds in res.get('dados_sql_list', []):
                    st.markdown(f"""
**{ds.get('nome_script', '?')}**

| Campo | Valor |
|---|---|
| Tabelas SQL | {', '.join(ds.get('tabelas', [])) if ds.get('tabelas') else 'N/A'} |
| Arquivo saída | `{ds.get('arquivo_saida') or 'N/A'}` |
| Operações | {descrever_operacao_sql(ds)} |
""")

with col_output:
    st.markdown("### 📊 Resultado")

    if st.session_state.resultados:
        for i, res in enumerate(st.session_state.resultados):
            cod = res['cod']
            svgs = res['svgs']

            if len(st.session_state.resultados) > 1:
                st.markdown(f"#### 🔹 {cod} — {res['nome_script']}")

            # ─── Downloads ──────────────────────────────
            st.markdown("##### 📥 Downloads")
            dl_cols = st.columns(3)

            with dl_cols[0]:
                st.download_button(
                    label=f"⬇ Característica .docx",
                    data=res['docx_bytes'],
                    file_name=f"Caracteristica_{cod}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_docx_{cod}_{i}",
                    use_container_width=True,
                )

            with dl_cols[1]:
                if res.get('fluxo_docx_bytes'):
                    st.download_button(
                        label=f"⬇ Fluxograma .docx",
                        data=res['fluxo_docx_bytes'],
                        file_name=f"Fluxograma_{cod}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_fluxo_docx_{cod}_{i}",
                        use_container_width=True,
                    )

            with dl_cols[2]:
                if len(svgs) == 1:
                    st.download_button(
                        label=f"⬇ Fluxograma .svg",
                        data=svgs[0].encode('utf-8'),
                        file_name=f"Fluxograma_{cod}.svg",
                        mime="image/svg+xml",
                        key=f"dl_svg_{cod}_{i}",
                        use_container_width=True,
                    )
                else:
                    # Combina múltiplas folhas em um SVG
                    W = 700
                    combined_els = []
                    total_h = 0
                    gap = 30
                    for svg_str in svgs:
                        vb = re.search(r'viewBox="0 0 (\d+) (\d+)"', svg_str)
                        h = int(vb.group(2)) if vb else 500
                        inner = re.search(r'<svg[^>]*>(.*)</svg>', svg_str, re.DOTALL)
                        if inner:
                            combined_els.append(f'<g transform="translate(0,{total_h})">{inner.group(1)}</g>')
                        total_h += h + gap
                    combined_svg = f'<svg viewBox="0 0 {W} {total_h}" xmlns="http://www.w3.org/2000/svg" width="{W}" style="background:#fff;max-width:100%">\n' + '\n'.join(combined_els) + '\n</svg>'
                    st.download_button(
                        label=f"⬇ Fluxograma_{cod} ({len(svgs)} folhas).svg",
                        data=combined_svg.encode('utf-8'),
                        file_name=f"Fluxograma_{cod}_Completo.svg",
                        mime="image/svg+xml",
                        key=f"dl_svg_{cod}_{i}",
                        use_container_width=True,
                    )

            # ─── Preview do fluxograma ──────────────────
            st.markdown("##### 👁 Preview do Fluxograma")
            for i, svg in enumerate(svgs):
                if len(svgs) > 1:
                    st.caption(f"Folha {i+1}/{len(svgs)}")
                st.markdown(svg, unsafe_allow_html=True)

            if len(st.session_state.resultados) > 1:
                st.markdown("---")

    else:
        st.markdown("""
        <div style="display:flex;flex-direction:column;align-items:center;justify-content:center;height:400px;color:#999;text-align:center">
            <div style="font-size:48px;margin-bottom:12px;opacity:.3">📊</div>
            <div style="font-size:14px;line-height:1.6">
                Faça upload do(s) script(s) SQL,<br>
                preencha os dados da rotina<br>
                e clique em <strong>Gerar Característica + Fluxograma</strong>
            </div>
        </div>
        """, unsafe_allow_html=True)
